import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
import logging

logger = logging.getLogger("app.tools")

def generate_docx(case_title: str, content_markdown: str, output_path: str) -> str:
    """
    Generates a DOCX document representing the drafted legal claim or complaint.
    """
    try:
        doc = Document()
        doc.add_heading(case_title, level=0)
        
        # Simple markdown-to-paragraph mapping
        for line in content_markdown.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        # Make sure directory path exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        logger.info(f"DOCX document successfully saved to: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        raise e

def generate_pdf(case_title: str, content_markdown: str, output_path: str) -> str:
    """
    Generates a clean, professional PDF document for courts/arbitrators using ReportLab.
    """
    try:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Setup ReportLab DocTemplate
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54, leftMargin=54,
            topMargin=54, bottomMargin=54
        )

        styles = getSampleStyleSheet()
        story = []

        # Create custom document styles
        title_style = ParagraphStyle(
            name='CaseTitleStyle',
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            alignment=1,  # Centered
            spaceAfter=20
        )
        
        heading_style = ParagraphStyle(
            name='CaseHeadingStyle',
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            spaceBefore=12,
            spaceAfter=6
        )

        body_style = ParagraphStyle(
            name='CaseBodyStyle',
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            spaceAfter=8
        )

        # Title block
        story.append(Paragraph(case_title, title_style))
        story.append(Spacer(1, 12))

        # Compile body paragraphs
        for line in content_markdown.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# ") or line.startswith("## ") or line.startswith("### "):
                clean_text = line.replace("#", "").strip()
                story.append(Paragraph(clean_text, heading_style))
            elif line.startswith("- ") or line.startswith("* "):
                clean_text = line[2:].strip()
                story.append(Paragraph(f"&bull; {clean_text}", body_style))
            else:
                story.append(Paragraph(line, body_style))

        # Build PDF file
        doc.build(story)
        logger.info(f"PDF document successfully saved to: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        raise e
